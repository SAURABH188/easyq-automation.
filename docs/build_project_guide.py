from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "EasyQ_Automation_Project_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "555555"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D3DF"


def color(hex_value):
    return RGBColor.from_string(hex_value)


def set_run(run, name="Calibri", size=None, color_hex=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color_hex is not None:
        run.font.color.rgb = color(color_hex)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def para_space(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, border_color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)


def cell_text(cell, text, bold=False, fill=None, size=9, color_hex="000000"):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    para_space(p, after=0, line=1.1)
    run = p.add_run(str(text))
    set_run(run, size=size, color_hex=color_hex, bold=bold)


def table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    borders(tbl)
    for i, header in enumerate(headers):
        tbl.rows[0].cells[i].width = Inches(widths[i])
        cell_text(tbl.rows[0].cells[i], header, bold=True, fill=header_fill, size=9.3, color_hex=INK)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            cell_text(cells[i], value, size=8.7)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return tbl


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    para_space(p)
    run = p.add_run(text)
    set_run(run)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        para_space(p, after=4)
        run = p.add_run(item)
        set_run(run, size=10.5)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        para_space(p, after=4)
        run = p.add_run(item)
        set_run(run, size=10.5)


def note(doc, title, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(tbl, border_color="B7C9DD")
    c = tbl.cell(0, 0)
    shade(c, "F4F8FC")
    margins(c, top=120, bottom=120, start=160, end=160)
    p = c.paragraphs[0]
    para_space(p, after=0, line=1.15)
    r = p.add_run(title + ": ")
    set_run(r, bold=True, color_hex=DARK_BLUE)
    r = p.add_run(text)
    set_run(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def code(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(tbl, border_color="DADCE0", size="4")
    c = tbl.cell(0, 0)
    shade(c, "F7F9FC")
    margins(c, top=120, bottom=120, start=160, end=160)
    p = c.paragraphs[0]
    para_space(p, after=0, line=1.05)
    for index, line in enumerate(text.strip("\n").splitlines()):
        if index:
            p.add_run("\n")
        r = p.add_run(line)
        set_run(r, name="Consolas", size=8.2, color_hex="222222")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def config_values():
    values = {}
    path = ROOT / "src/test/resources/config.properties"
    for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        values[key.strip()] = value.strip()
    return values


def pom_properties():
    text = (ROOT / "pom.xml").read_text(encoding="utf-8", errors="ignore")
    return dict(re.findall(r"<([^/][^>]+)>([^<>]+)</\1>", text))


def suites_by_class():
    mapping = {}
    for xml in ROOT.glob("testng*.xml"):
        text = xml.read_text(encoding="utf-8", errors="ignore")
        for cls in re.findall(r'<class\s+name="tests\.([^"]+)"', text):
            mapping.setdefault(cls, []).append(xml.name)
    return mapping


def test_modules():
    friendly = {
        "EasyQActiveLogsTest": ("Active Logs", "Logs, grouping, date/time, scrolling and pagination."),
        "EasyQCapaTest": ("CAPA", "Access, draft/review/approval, lifecycle and role restrictions."),
        "EasyQComplaintManagementTest": ("Complaint Management", "Complaint form, save/submit, upload, CAPA linkage, filters and actions."),
        "EasyQCredentialValidationTest": ("Credential Validation", "Configured role users can authenticate."),
        "EasyQDashboardDetailsTest": ("Dashboard Details", "Widget counts, navigation, severity, role-based and null/large data."),
        "EasyQDashboardTest": ("Dashboard", "Dashboard load, reload, widgets, toggles and basic layout."),
        "EasyQDocumentManagementTest": ("Document Management", "Upload, drafts, review/approval, download, versions and obsolete flow."),
        "EasyQDynamicWorkflowTest": ("Dynamic Workflow", "State-aware workflows for dynamic UI conditions."),
        "EasyQHomepage": ("Homepage / Dashboard Smoke", "Dashboard card visibility smoke checks."),
        "EasyQLoginPasswordRemainingTest": ("Login + Forgot Password Remaining", "Additional login, validation, browser and forgot-password scenarios."),
        "EasyQLoginUiUxComprehensiveTest": ("Login UI/UX Comprehensive", "Login UI, validation, keyboard, responsive and edge cases."),
        "EasyQLoginValidation": ("Login Validation", "Manual login IDs, valid/invalid login and forgot-password navigation."),
        "EasyQManagementReviewTest": ("Management Review", "MR assignment, schedule, MOM, approval and task flow."),
        "EasyQNotificationsTest": ("Notifications", "Panel, read/unread, ordering, show more, redirection and data."),
        "EasyQProductsTest": ("Products", "List, search, filters, status, delete, duplicates and text handling."),
        "EasyQQualityObjectiveTest": ("Quality Objective", "Initiation, draft, review, approval, versions and access checks."),
        "EasyQQualityPolicyTest": ("Quality Policy", "Policy draft/review/approval, active/inactive and versioning."),
        "EasyQResponsibilityAuthorityTest": ("Responsibility & Authority", "Rows, drafts, review/approval, downloads and access."),
        "EasyQTrainingTest": ("Training", "Assignments, document linkage, completion, counts, roles and dynamic states."),
        "EasyQUserManagementTest": ("User Management", "Users, licenses, tabs, add/edit/disable, roles, access and cleanup."),
        "LoginTest": ("Login Smoke", "Basic login page and valid login smoke checks."),
    }
    suite_map = suites_by_class()
    rows = []
    for java in sorted((ROOT / "src/test/java/tests").glob("*.java")):
        name = java.stem
        module, focus = friendly.get(name, (name, "Module-specific UI and workflow checks."))
        count = len(re.findall(r"@Test\b", java.read_text(encoding="utf-8", errors="ignore")))
        rows.append((module, java.name, ", ".join(suite_map.get(name, ["-"])), str(count), focus))
    return rows


cfg = config_values()
props = pom_properties()
modules = test_modules()
total_tests = sum(int(row[3]) for row in modules)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color_hex in [
    ("Heading 1", 16, BLUE),
    ("Heading 2", 13, BLUE),
    ("Heading 3", 12, DARK_BLUE),
]:
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color(color_hex)
    style.font.bold = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("EasyQ Automation Project Guide"), size=9, color_hex=MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer.add_run("Internal QA automation reference - generated July 1, 2026"), size=8.5, color_hex=MUTED)

p = doc.add_paragraph()
para_space(p, before=72, after=6)
set_run(p.add_run("EasyQ Automation Project Guide"), size=28, color_hex=INK, bold=True)
p = doc.add_paragraph()
para_space(p, after=18)
set_run(p.add_run("Selenium Java TestNG framework for EasyQ beta environment"), size=14, color_hex=MUTED)

table(doc, ["Field", "Value"], [
    ("Project folder", str(ROOT)),
    ("Beta URL", cfg.get("baseUrl", "")),
    ("Repository", "https://github.com/SAURABH188/easyq-automation..git"),
    ("Branch", "master"),
    ("Prepared for", "QA lead self-use and demo walkthrough"),
    ("Document date", "July 1, 2026"),
], [1.6, 4.8], header_fill=LIGHT_GRAY)
note(doc, "Security note", "This guide intentionally does not include real account passwords. Store passwords in Eclipse environment variables or secrets.local.properties, and never commit them to GitHub.")
body(doc, "Purpose: This guide explains how the EasyQ automation project is structured, how to set it up, how scripts are written, how to run module suites, and how to present the framework during a demo.")

doc.add_page_break()
heading(doc, "Table of Contents", 1)
bullets(doc, [
    "1. Project Snapshot",
    "2. Framework and Tool Stack",
    "3. Project Folder Structure",
    "4. Configuration and Environment Variables",
    "5. Step-by-Step Setup in Eclipse",
    "6. How Scripts Are Written",
    "7. Wait Strategy and Stability Handling",
    "8. TestNG Suites and Module Coverage",
    "9. Running Tests",
    "10. Dynamic Workflow and Role-Based Testing",
    "11. Reporting and Debugging",
    "12. GitHub Workflow",
    "13. Demo Talking Points",
    "14. Maintenance Checklist",
    "Appendix A. Useful Commands",
    "Appendix B. Current Suite Map",
])

doc.add_page_break()
heading(doc, "1. Project Snapshot", 1)
body(doc, "EasyQ automation is a Selenium WebDriver framework built in Java for validating the EasyQ beta web application. It uses TestNG for test organization, Maven for dependency management, WebDriverManager for browser driver setup, and Eclipse as the primary IDE.")
table(doc, ["Area", "Current Details"], [
    ("Application under test", "EasyQ beta web app"),
    ("Environment", cfg.get("baseUrl", "")),
    ("Primary browser", cfg.get("browser", "chrome")),
    ("Automation style", "Selenium WebDriver + TestNG module suites"),
    ("Current test classes", str(len(modules))),
    ("Current @Test methods detected", str(total_tests)),
    ("Failure artifacts", "Console failure reason and screenshots under test-output/screenshots"),
], [2.1, 4.3])
heading(doc, "Automation Goals", 2)
bullets(doc, [
    "Validate login, dashboard and major QMS modules in beta.",
    "Map automation methods back to manual test case IDs where available.",
    "Use role-based users for Admin, Document Controller and Assignee behavior.",
    "Handle dynamic screens by checking real UI states instead of blindly skipping workflows.",
    "Generate useful failure details for fast debugging during demo and local execution.",
])

heading(doc, "2. Framework and Tool Stack", 1)
table(doc, ["Tool", "Purpose", "Current Version / Source"], [
    ("Java", "Programming language used for Selenium scripts.", props.get("maven.compiler.release", "17")),
    ("Maven", "Dependency management, compile lifecycle and suite execution.", "pom.xml"),
    ("Selenium WebDriver", "Browser automation and UI interaction.", props.get("selenium.version", "4.33.0")),
    ("TestNG", "Test annotations, priorities, suites, listeners and reports.", props.get("testng.version", "7.8.0")),
    ("WebDriverManager", "Automatically resolves browser drivers.", props.get("webdrivermanager.version", "6.1.0")),
    ("SLF4J Simple", "Logging provider to avoid missing logger warnings.", props.get("slf4j.version", "1.7.36")),
    ("Eclipse IDE", "Primary authoring and run environment.", "Run as TestNG Suite"),
    ("GitHub", "Source control and sharing.", "origin remote configured"),
], [1.55, 3.25, 1.6])
body(doc, "The Maven project keeps framework versions in pom.xml properties. The surefire plugin can run any suite file by changing the suiteXmlFile property.")

heading(doc, "3. Project Folder Structure", 1)
table(doc, ["Path / Pattern", "Purpose"], [
    ("src/test/java/base", "Base test setup, browser creation and shared start/stop flow."),
    ("src/test/java/pages", "Page Object classes, currently including LoginPage."),
    ("src/test/java/tests", "TestNG test classes for login, dashboard and QMS modules."),
    ("src/test/java/utils", "Reusable helpers for waits, actions, dynamic workflows, config and screenshots."),
    ("src/test/resources", "Framework config and optional local secrets file support."),
    ("testng*.xml", "Suite files for full run and module-wise execution."),
    ("test-output/screenshots", "Failure screenshots generated by TestListener."),
], [2.4, 4.0])
code(doc, """easyq-automation/
  pom.xml
  testng.xml
  testng-*.xml
  src/test/java/base/BaseTest.java
  src/test/java/pages/LoginPage.java
  src/test/java/tests/*.java
  src/test/java/utils/*.java
  src/test/resources/config.properties
  test-output/screenshots/""")

heading(doc, "4. Configuration and Environment Variables", 1)
heading(doc, "config.properties", 2)
table(doc, ["Key", "Current Value / Use"], [
    ("baseUrl", cfg.get("baseUrl", "")),
    ("browser", cfg.get("browser", "")),
    ("explicitWait", cfg.get("explicitWait", "")),
    ("actionDelayMs", cfg.get("actionDelayMs", "")),
    ("pageSettleMs", cfg.get("pageSettleMs", "")),
    ("highlightActions", cfg.get("highlightActions", "")),
    ("allowWorkflowMutations", cfg.get("allowWorkflowMutations", "")),
    ("allowPasswordResetMutation", cfg.get("allowPasswordResetMutation", "")),
], [2.1, 4.3])
body(doc, "Passwords should be entered in Eclipse Run Configuration > Environment, as system properties, or in a local secrets.local.properties file that is not committed.")
table(doc, ["Variable", "Use"], [
    ("EASYQ_ADMIN_PASSWORD or EASYQ_PASSWORD", "Admin password used for valid login scenarios."),
    ("EASYQ_DOC_CONTROLLER_PASSWORD", "Document Controller role password."),
    ("EASYQ_ASSIGNEE_SWATI_PASSWORD", "Assignee role password for role-restriction checks."),
    ("EASYQ_TEST_USER_*", "Disposable automation user data for User Management workflows."),
    ("EASYQ_ALLOW_WORKFLOW_MUTATIONS", "Allows create/edit/delete workflow tests when true."),
    ("EASYQ_VISUAL_DELAY_MS", "Override for visible action delay while watching automation run."),
    ("EASYQ_PAGE_SETTLE_MS", "Override for post-load DOM settle delay."),
], [2.55, 3.85])
note(doc, "Recommended demo setting", "Use EASYQ_VISUAL_DELAY_MS=1500 or 2000 when showing automation to the team, so actions are visible.")

heading(doc, "5. Step-by-Step Setup in Eclipse", 1)
numbers(doc, [
    "Install or select JDK 17 in Eclipse. Confirm Project > Properties > Java Compiler uses Java 17.",
    "Open Eclipse and choose File > Import > Maven > Existing Maven Projects.",
    "Select the easyq-automation folder under the Codex output directory.",
    "Let Maven download dependencies. If the Maven cache is locked, close Eclipse/Java processes and retry Maven Update.",
    "Install or enable the TestNG plugin in Eclipse if TestNG run options are not visible.",
    "Open Run Configurations > TestNG > your suite configuration.",
    "Set Environment variables such as EASYQ_ADMIN_PASSWORD, role passwords and optional visual delay.",
    "Run testng-user-management.xml or any module XML as a TestNG Suite.",
    "Review Eclipse Console output and screenshots in test-output/screenshots when failures occur.",
])
heading(doc, "Eclipse Environment Example", 2)
code(doc, """EASYQ_ADMIN_PASSWORD=<admin password>
EASYQ_DOC_CONTROLLER_PASSWORD=<doc controller password>
EASYQ_ASSIGNEE_SWATI_PASSWORD=<assignee password>
EASYQ_VISUAL_DELAY_MS=1500
EASYQ_PAGE_SETTLE_MS=1200""")

heading(doc, "6. How Scripts Are Written", 1)
body(doc, "Each script is a TestNG Java class. Test methods are marked with @Test, usually with priority and description. Manual test case IDs are placed near the script as normal comments where available.")
heading(doc, "Typical Test Flow", 2)
numbers(doc, [
    "Start Chrome using WebDriverManager.",
    "Open the EasyQ beta login page.",
    "Wait for page load and visible fields.",
    "Login using configured username and password.",
    "Navigate to the target module by sidebar, profile dropdown or route fallback.",
    "Wait for the module to render and loaders to finish.",
    "Perform the action or assertion for the manual test scenario.",
    "On failure, print the failure reason and capture a screenshot.",
])
heading(doc, "Script Pattern Example", 2)
code(doc, """@Test(priority = 1, description = \"Verify User Management page loads successfully\")
// Manual Test Case ID: TC279
public void verifyUserManagementPageLoadsSuccessfully() {
    Assert.assertTrue(
        wait.until(ExpectedConditions.visibilityOfElementLocated(userManagementTitle)).isDisplayed(),
        \"User Management page should load successfully\"
    );
}""")
heading(doc, "Locator Strategy", 2)
bullets(doc, [
    "Use flexible XPath/CSS selectors when stable IDs are not available in beta UI.",
    "Prefer visible text, aria-label, title, role and formcontrolname over brittle absolute XPath.",
    "Future improvement: ask developers to add stable data-testid attributes for important controls.",
])

heading(doc, "7. Wait Strategy and Stability Handling", 1)
body(doc, "The beta environment can load slowly. The framework combines explicit waits, app-load waits, visible loader checks, DOM settle delay and action observation delay.")
table(doc, ["Layer", "Behavior"], [
    ("BaseTest", "Sets browser, page-load timeout, script timeout, opens base URL and waits for app load."),
    ("WaitHelper", "Waits for document.readyState, Angular stability, loader disappearance and DOM settle."),
    ("ActionHelper", "Waits before click/type, scrolls/highlights elements and pauses after each action."),
    ("TestNG XML", "configfailurepolicy=continue helps avoid a skip cascade after setup/browser issues."),
    ("User Management special handling", "If Chrome is manually closed or setup fails, the test is failed clearly instead of silently skipped."),
], [2.0, 4.4])
note(doc, "Why this matters", "If scripts run faster than the platform can render, Selenium looks for elements too early and the test breaks.")

heading(doc, "8. TestNG Suites and Module Coverage", 1)
body(doc, f"The project currently contains {len(modules)} Java test classes and {total_tests} detected @Test methods. Module suite XML files are available so you can run one area at a time.")
table(doc, ["Module", "Class", "Suite XML", "Tests", "Focus"], modules, [1.25, 1.65, 1.55, 0.45, 1.6])

heading(doc, "9. Running Tests", 1)
heading(doc, "Run from Eclipse", 2)
numbers(doc, [
    "Right-click the suite file, for example testng-user-management.xml.",
    "Choose Run As > TestNG Suite.",
    "Watch the browser actions and Eclipse Console logs.",
    "If one test fails, open the failure reason and screenshot path printed in the console.",
    "For failed-only reruns, use Eclipse TestNG failed suite output when available, or run the specific method/class.",
])
heading(doc, "Run from Maven", 2)
code(doc, """mvn test
mvn test -DsuiteXmlFile=testng-user-management.xml
mvn test -DsuiteXmlFile=testng-login-ui-ux.xml""")
note(doc, "Maven cache issue", "If Maven reports AccessDeniedException under .m2, close Eclipse/Java processes and remove the affected plugin folder only after verifying the path.")

heading(doc, "10. Dynamic Workflow and Role-Based Testing", 1)
body(doc, "Some EasyQ workflows are dynamic: a module may show existing data, no data, restricted access, draft state, review state or approved state. The framework checks valid states instead of directly skipping workflows.")
bullets(doc, [
    "DynamicWorkflowHelper checks for meaningful page states such as Draft, Review, Approved, Pending, Completed, Closed, Open, No Data or Access Denied.",
    "Mutation-heavy tests are controlled by allowWorkflowMutations or EASYQ_ALLOW_WORKFLOW_MUTATIONS.",
    "Role-based tests use configured Admin, Document Controller and Assignee accounts.",
    "Disposable user tests use EASYQ_TEST_USER_* variables and include cleanup behavior.",
])
table(doc, ["Role Type", "Username Key", "Password Key"], [
    ("Admin", "EASYQ_ADMIN_USERNAME", "EASYQ_ADMIN_PASSWORD or EASYQ_PASSWORD"),
    ("Document Controller", "EASYQ_DOC_CONTROLLER_USERNAME", "EASYQ_DOC_CONTROLLER_PASSWORD"),
    ("Assignee", "EASYQ_ASSIGNEE_SWATI_USERNAME", "EASYQ_ASSIGNEE_SWATI_PASSWORD"),
    ("Other Assignees", "EASYQ_ASSIGNEE_*_USERNAME", "matching password variable if used"),
], [1.6, 2.3, 2.5])

heading(doc, "11. Reporting and Debugging", 1)
body(doc, "TestListener prints failed test name, test data, failure reason and screenshot path. Several tests also print navigation or login step logs to make stuck points visible.")
code(doc, """FAILED TEST: verifyUserManagementPageLoadsSuccessfully
FAILURE REASON: <root cause>
SCREENSHOT: C:\\...\\test-output\\screenshots\\verifyUserManagementPageLoadsSuccessfully.png""")
heading(doc, "Common Issues and Fixes", 2)
table(doc, ["Issue", "Meaning", "Fix"], [
    ("EASYQ_PASSWORD is required", "Password variable is missing.", "Set EASYQ_ADMIN_PASSWORD or EASYQ_PASSWORD in Eclipse."),
    ("Login stays on login page", "Password may be wrong, account inactive, or field/click issue.", "Verify credentials manually, then rerun."),
    ("CDP implementation warning", "Chrome is newer than Selenium devtools package.", "Usually non-blocking unless DevTools APIs are used."),
    ("NoSuchWindowException", "Chrome was manually closed or crashed.", "Rerun; teardown is now safer and next test starts fresh."),
    ("Timeout waiting for element", "Page, module or control did not render in time.", "Use module suite, confirm navigation, increase visual delay/wait."),
    ("Unsupported class file major version", "Java/Maven/TestNG runtime mismatch.", "Use JDK 17 consistently in Eclipse and Maven."),
    (".m2 AccessDeniedException", "Maven cache folder locked or permission issue.", "Close Java/Eclipse processes; remove only affected plugin cache if it exists."),
], [1.65, 2.35, 2.4])

heading(doc, "12. GitHub Workflow", 1)
body(doc, "The project is connected to GitHub with origin remote configured. Use small commits after stable changes, and avoid committing secrets.local.properties, screenshots containing sensitive data, or real passwords.")
table(doc, ["Item", "Value"], [
    ("Remote", "https://github.com/SAURABH188/easyq-automation..git"),
    ("Current branch", "master"),
    ("Recommended commit scope", "One framework/module change at a time"),
], [1.8, 4.6])
heading(doc, "Manual Push Steps from Eclipse", 2)
numbers(doc, [
    "Right-click project > Team > Commit.",
    "Review changed files and enter a clear commit message.",
    "Click Commit and Push, or Commit first and push later.",
    "If Push Ref Specification appears empty, choose Add All Branches Spec or select master -> master.",
    "Confirm Push Results shows master -> master or your current branch pushed successfully.",
])
heading(doc, "Command Line Backup", 2)
code(doc, """git status
git add .
git commit -m \"Update EasyQ automation framework waits and documentation\"
git push origin master""")

heading(doc, "13. Demo Talking Points", 1)
numbers(doc, [
    "Start by explaining the purpose: automate EasyQ beta validation across login, dashboard and QMS modules.",
    "Show pom.xml and explain Java, Selenium, TestNG and WebDriverManager.",
    "Show config.properties and explain that passwords are outside source code.",
    "Show testng XML files and explain module-wise execution.",
    "Open one test class and point out @Test description and manual test case ID comments.",
    "Show BaseTest, ActionHelper and WaitHelper to explain stability handling.",
    "Run a small suite such as login UI/UX or User Management page load.",
    "Show console output and screenshot behavior if a test fails.",
    "Close with next improvements: stable locators, CI, reports and cleaner page objects.",
])

heading(doc, "14. Maintenance Checklist", 1)
table(doc, ["Checklist Item", "Why It Matters"], [
    ("Keep credentials outside GitHub.", "Prevents password leakage."),
    ("Use module XML while debugging.", "Faster feedback than running the full suite."),
    ("Prefer stable locators.", "Reduces failures caused by UI layout changes."),
    ("Add manual test case ID comments.", "Keeps automation linked with QA test cases."),
    ("Capture useful assertion messages.", "Makes failed tests understandable without reopening code."),
    ("Use dynamic state handling for workflows.", "Avoids false failures when application data state changes."),
    ("Review screenshots after failures.", "Fastest way to confirm script, data, role or app-loading issue."),
    ("Commit after a passing compile.", "Keeps GitHub history clean and demo-safe."),
], [2.25, 4.15])

heading(doc, "Appendix A. Useful Commands", 1)
code(doc, """# Run full suite if Maven cache is healthy
mvn test

# Run a module suite
mvn test -DsuiteXmlFile=testng-user-management.xml

# Check Git status
git status

# Push current branch
git push origin master

# Useful Eclipse run environment examples
EASYQ_ADMIN_PASSWORD=<password>
EASYQ_VISUAL_DELAY_MS=1500
EASYQ_PAGE_SETTLE_MS=1200""")

heading(doc, "Appendix B. Current Suite Map", 1)
suite_rows = []
for xml in sorted(ROOT.glob("testng*.xml")):
    text = xml.read_text(encoding="utf-8", errors="ignore")
    suite = re.search(r'<suite\s+name="([^"]+)"', text)
    classes = ", ".join(re.findall(r'<class\s+name="tests\.([^"]+)"', text)) or "-"
    suite_rows.append((xml.name, suite.group(1) if suite else "-", classes))
table(doc, ["Suite File", "Suite Name", "Classes"], suite_rows, [1.65, 2.2, 2.55])

heading(doc, "Appendix C. Recommended Next Improvements", 1)
bullets(doc, [
    "Ask developers to add data-testid attributes to key buttons, inputs, menus and module links.",
    "Move remaining module classes toward a common BaseTest and page-object structure.",
    "Add ExtentReports or Allure for richer demo reporting.",
    "Create CI execution for nightly smoke suites once credentials are stored in CI secrets.",
    "Keep one stable smoke suite separate from long workflow mutation tests.",
])

core_props = doc.core_properties
core_props.title = "EasyQ Automation Project Guide"
core_props.subject = "Selenium Java TestNG framework documentation"
core_props.author = "Codex for Saurabh"
core_props.comments = "Internal QA reference; no passwords included."

doc.save(OUT)


def build_pdf():
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
        KeepTogether,
    )

    pdf_path = ROOT / "docs" / "EasyQ_Automation_Project_Guide.pdf"
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12.5,
        spaceAfter=6,
        textColor=colors.HexColor("#111111"),
    )
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=29,
        textColor=colors.HexColor("#0B2545"),
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#555555"),
        spaceAfter=18,
    )
    h1_style = ParagraphStyle(
        "Heading1Custom",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12,
        spaceAfter=8,
    )
    h2_style = ParagraphStyle(
        "Heading2Custom",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=14,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=9,
        spaceAfter=5,
    )
    code_style = ParagraphStyle(
        "Code",
        parent=styles["BodyText"],
        fontName="Courier",
        fontSize=7.6,
        leading=9.2,
        textColor=colors.HexColor("#222222"),
    )
    small_style = ParagraphStyle(
        "Small",
        parent=body_style,
        fontSize=8,
        leading=10,
    )
    note_style = ParagraphStyle(
        "Note",
        parent=body_style,
        backColor=colors.HexColor("#F4F8FC"),
        borderColor=colors.HexColor("#B7C9DD"),
        borderWidth=0.5,
        borderPadding=7,
        spaceBefore=4,
        spaceAfter=8,
    )

    def esc(value):
        return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def p(text, style=body_style):
        return Paragraph(esc(text), style)

    def h1(text):
        return Paragraph(esc(text), h1_style)

    def h2(text):
        return Paragraph(esc(text), h2_style)

    def bullet_list(items):
        out = []
        for item in items:
            out.append(Paragraph("&#8226; " + esc(item), body_style))
        return out

    def number_list(items):
        out = []
        for idx, item in enumerate(items, start=1):
            out.append(Paragraph(f"{idx}. " + esc(item), body_style))
        return out

    def rl_table(headers, rows, widths, font_size=7.7):
        header_style = ParagraphStyle(
            "TableHeader",
            parent=small_style,
            fontName="Helvetica-Bold",
            fontSize=font_size,
            leading=font_size + 2,
            textColor=colors.HexColor("#0B2545"),
        )
        cell_style = ParagraphStyle(
            "TableCell",
            parent=small_style,
            fontName="Helvetica",
            fontSize=font_size,
            leading=font_size + 2,
        )
        data = [[Paragraph(esc(h), header_style) for h in headers]]
        for row in rows:
            data.append([Paragraph(esc(value), cell_style) for value in row])
        tbl = Table(data, colWidths=[w * inch for w in widths], repeatRows=1, hAlign="CENTER")
        tbl.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D3DF")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return tbl

    def code_block(text):
        lines = esc(text.strip("\n")).replace("\n", "<br/>")
        return Table(
            [[Paragraph(lines, code_style)]],
            colWidths=[6.35 * inch],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FC")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#DADCE0")),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]),
        )

    def on_page(canvas, pdf_doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(0.72 * inch, 0.5 * inch, "EasyQ Automation Project Guide")
        canvas.drawRightString(7.78 * inch, 0.5 * inch, f"Page {pdf_doc.page}")
        canvas.restoreState()

    story = []
    story.append(Paragraph("EasyQ Automation Project Guide", title_style))
    story.append(Paragraph("Selenium Java TestNG framework for EasyQ beta environment", subtitle_style))
    story.append(rl_table(["Field", "Value"], [
        ("Project folder", str(ROOT)),
        ("Beta URL", cfg.get("baseUrl", "")),
        ("Repository", "https://github.com/SAURABH188/easyq-automation..git"),
        ("Branch", "master"),
        ("Prepared for", "QA lead self-use and demo walkthrough"),
        ("Document date", "July 1, 2026"),
    ], [1.55, 4.8], font_size=8))
    story.append(Spacer(1, 8))
    story.append(Paragraph("<b>Security note:</b> This guide intentionally does not include real account passwords. Store passwords in Eclipse environment variables or secrets.local.properties, and never commit them to GitHub.", note_style))
    story.append(p("Purpose: This guide explains how the EasyQ automation project is structured, how to set it up, how scripts are written, how to run module suites, and how to present the framework during a demo."))
    story.append(PageBreak())

    story.append(h1("Table of Contents"))
    story.extend(bullet_list([
        "Project Snapshot",
        "Framework and Tool Stack",
        "Project Folder Structure",
        "Configuration and Environment Variables",
        "Step-by-Step Setup in Eclipse",
        "How Scripts Are Written",
        "Wait Strategy and Stability Handling",
        "TestNG Suites and Module Coverage",
        "Running Tests",
        "Dynamic Workflow and Role-Based Testing",
        "Reporting and Debugging",
        "GitHub Workflow",
        "Demo Talking Points",
        "Maintenance Checklist",
        "Appendix A. Useful Commands",
        "Appendix B. Current Suite Map",
    ]))
    story.append(PageBreak())

    story.append(h1("1. Project Snapshot"))
    story.append(p("EasyQ automation is a Selenium WebDriver framework built in Java for validating the EasyQ beta web application. It uses TestNG for test organization, Maven for dependency management, WebDriverManager for browser driver setup, and Eclipse as the primary IDE."))
    story.append(rl_table(["Area", "Current Details"], [
        ("Application under test", "EasyQ beta web app"),
        ("Environment", cfg.get("baseUrl", "")),
        ("Primary browser", cfg.get("browser", "chrome")),
        ("Automation style", "Selenium WebDriver + TestNG module suites"),
        ("Current test classes", str(len(modules))),
        ("Current @Test methods detected", str(total_tests)),
        ("Failure artifacts", "Console failure reason and screenshots under test-output/screenshots"),
    ], [2.0, 4.35], font_size=8))
    story.append(h2("Automation Goals"))
    story.extend(bullet_list([
        "Validate login, dashboard and major QMS modules in beta.",
        "Map automation methods back to manual test case IDs where available.",
        "Use role-based users for Admin, Document Controller and Assignee behavior.",
        "Handle dynamic screens by checking real UI states instead of blindly skipping workflows.",
        "Generate useful failure details for fast debugging during demo and local execution.",
    ]))

    story.append(h1("2. Framework and Tool Stack"))
    story.append(rl_table(["Tool", "Purpose", "Current Version / Source"], [
        ("Java", "Programming language used for Selenium scripts.", props.get("maven.compiler.release", "17")),
        ("Maven", "Dependency management, compile lifecycle and suite execution.", "pom.xml"),
        ("Selenium WebDriver", "Browser automation and UI interaction.", props.get("selenium.version", "4.33.0")),
        ("TestNG", "Test annotations, priorities, suites, listeners and reports.", props.get("testng.version", "7.8.0")),
        ("WebDriverManager", "Automatically resolves browser drivers.", props.get("webdrivermanager.version", "6.1.0")),
        ("SLF4J Simple", "Logging provider to avoid missing logger warnings.", props.get("slf4j.version", "1.7.36")),
        ("Eclipse IDE", "Primary authoring and run environment.", "Run as TestNG Suite"),
        ("GitHub", "Source control and sharing.", "origin remote configured"),
    ], [1.45, 3.25, 1.65], font_size=7.5))
    story.append(p("The Maven project keeps framework versions in pom.xml properties. The surefire plugin can run any suite file by changing the suiteXmlFile property."))

    story.append(h1("3. Project Folder Structure"))
    story.append(rl_table(["Path / Pattern", "Purpose"], [
        ("src/test/java/base", "Base test setup, browser creation and shared start/stop flow."),
        ("src/test/java/pages", "Page Object classes, currently including LoginPage."),
        ("src/test/java/tests", "TestNG test classes for login, dashboard and QMS modules."),
        ("src/test/java/utils", "Reusable helpers for waits, actions, dynamic workflows, config and screenshots."),
        ("src/test/resources", "Framework config and optional local secrets file support."),
        ("testng*.xml", "Suite files for full run and module-wise execution."),
        ("test-output/screenshots", "Failure screenshots generated by TestListener."),
    ], [2.2, 4.15], font_size=8))
    story.append(code_block("""easyq-automation/
  pom.xml
  testng.xml
  testng-*.xml
  src/test/java/base/BaseTest.java
  src/test/java/pages/LoginPage.java
  src/test/java/tests/*.java
  src/test/java/utils/*.java
  src/test/resources/config.properties
  test-output/screenshots/"""))

    story.append(h1("4. Configuration and Environment Variables"))
    story.append(h2("config.properties"))
    story.append(rl_table(["Key", "Current Value / Use"], [
        ("baseUrl", cfg.get("baseUrl", "")),
        ("browser", cfg.get("browser", "")),
        ("explicitWait", cfg.get("explicitWait", "")),
        ("actionDelayMs", cfg.get("actionDelayMs", "")),
        ("pageSettleMs", cfg.get("pageSettleMs", "")),
        ("highlightActions", cfg.get("highlightActions", "")),
        ("allowWorkflowMutations", cfg.get("allowWorkflowMutations", "")),
        ("allowPasswordResetMutation", cfg.get("allowPasswordResetMutation", "")),
    ], [2.05, 4.3], font_size=8))
    story.append(p("Passwords should be entered in Eclipse Run Configuration > Environment, as system properties, or in a local secrets.local.properties file that is not committed."))
    story.append(rl_table(["Variable", "Use"], [
        ("EASYQ_ADMIN_PASSWORD or EASYQ_PASSWORD", "Admin password used for valid login scenarios."),
        ("EASYQ_DOC_CONTROLLER_PASSWORD", "Document Controller role password."),
        ("EASYQ_ASSIGNEE_SWATI_PASSWORD", "Assignee role password for role-restriction checks."),
        ("EASYQ_TEST_USER_*", "Disposable automation user data for User Management workflows."),
        ("EASYQ_ALLOW_WORKFLOW_MUTATIONS", "Allows create/edit/delete workflow tests when true."),
        ("EASYQ_VISUAL_DELAY_MS", "Override for visible action delay while watching automation run."),
        ("EASYQ_PAGE_SETTLE_MS", "Override for post-load DOM settle delay."),
    ], [2.45, 3.9], font_size=7.6))
    story.append(Paragraph("<b>Recommended demo setting:</b> Use EASYQ_VISUAL_DELAY_MS=1500 or 2000 when showing automation to the team, so actions are visible.", note_style))

    story.append(h1("5. Step-by-Step Setup in Eclipse"))
    story.extend(number_list([
        "Install or select JDK 17 in Eclipse. Confirm Project > Properties > Java Compiler uses Java 17.",
        "Open Eclipse and choose File > Import > Maven > Existing Maven Projects.",
        "Select the easyq-automation folder under the Codex output directory.",
        "Let Maven download dependencies. If the Maven cache is locked, close Eclipse/Java processes and retry Maven Update.",
        "Install or enable the TestNG plugin in Eclipse if TestNG run options are not visible.",
        "Open Run Configurations > TestNG > your suite configuration.",
        "Set Environment variables such as EASYQ_ADMIN_PASSWORD, role passwords and optional visual delay.",
        "Run testng-user-management.xml or any module XML as a TestNG Suite.",
        "Review Eclipse Console output and screenshots in test-output/screenshots when failures occur.",
    ]))
    story.append(h2("Eclipse Environment Example"))
    story.append(code_block("""EASYQ_ADMIN_PASSWORD=<admin password>
EASYQ_DOC_CONTROLLER_PASSWORD=<doc controller password>
EASYQ_ASSIGNEE_SWATI_PASSWORD=<assignee password>
EASYQ_VISUAL_DELAY_MS=1500
EASYQ_PAGE_SETTLE_MS=1200"""))

    story.append(h1("6. How Scripts Are Written"))
    story.append(p("Each script is a TestNG Java class. Test methods are marked with @Test, usually with priority and description. Manual test case IDs are placed near the script as normal comments where available."))
    story.append(h2("Typical Test Flow"))
    story.extend(number_list([
        "Start Chrome using WebDriverManager.",
        "Open the EasyQ beta login page.",
        "Wait for page load and visible fields.",
        "Login using configured username and password.",
        "Navigate to the target module by sidebar, profile dropdown or route fallback.",
        "Wait for the module to render and loaders to finish.",
        "Perform the action or assertion for the manual test scenario.",
        "On failure, print the failure reason and capture a screenshot.",
    ]))
    story.append(h2("Script Pattern Example"))
    story.append(code_block("""@Test(priority = 1, description = \"Verify User Management page loads successfully\")
// Manual Test Case ID: TC279
public void verifyUserManagementPageLoadsSuccessfully() {
    Assert.assertTrue(
        wait.until(ExpectedConditions.visibilityOfElementLocated(userManagementTitle)).isDisplayed(),
        \"User Management page should load successfully\"
    );
}"""))
    story.append(h2("Locator Strategy"))
    story.extend(bullet_list([
        "Use flexible XPath/CSS selectors when stable IDs are not available in beta UI.",
        "Prefer visible text, aria-label, title, role and formcontrolname over brittle absolute XPath.",
        "Future improvement: ask developers to add stable data-testid attributes for important controls.",
    ]))

    story.append(h1("7. Wait Strategy and Stability Handling"))
    story.append(p("The beta environment can load slowly. The framework combines explicit waits, app-load waits, visible loader checks, DOM settle delay and action observation delay."))
    story.append(rl_table(["Layer", "Behavior"], [
        ("BaseTest", "Sets browser, page-load timeout, script timeout, opens base URL and waits for app load."),
        ("WaitHelper", "Waits for document.readyState, Angular stability, loader disappearance and DOM settle."),
        ("ActionHelper", "Waits before click/type, scrolls/highlights elements and pauses after each action."),
        ("TestNG XML", "configfailurepolicy=continue helps avoid a skip cascade after setup/browser issues."),
        ("User Management special handling", "If Chrome is manually closed or setup fails, the test is failed clearly instead of silently skipped."),
    ], [1.85, 4.5], font_size=7.8))
    story.append(Paragraph("<b>Why this matters:</b> If scripts run faster than the platform can render, Selenium looks for elements too early and the test breaks.", note_style))

    story.append(h1("8. TestNG Suites and Module Coverage"))
    story.append(p(f"The project currently contains {len(modules)} Java test classes and {total_tests} detected @Test methods. Module suite XML files are available so you can run one area at a time."))
    story.append(rl_table(["Module", "Class", "Suite XML", "Tests", "Focus"], modules, [1.1, 1.45, 1.45, 0.4, 1.95], font_size=6.4))

    story.append(h1("9. Running Tests"))
    story.append(h2("Run from Eclipse"))
    story.extend(number_list([
        "Right-click the suite file, for example testng-user-management.xml.",
        "Choose Run As > TestNG Suite.",
        "Watch the browser actions and Eclipse Console logs.",
        "If one test fails, open the failure reason and screenshot path printed in the console.",
        "For failed-only reruns, use Eclipse TestNG failed suite output when available, or run the specific method/class.",
    ]))
    story.append(h2("Run from Maven"))
    story.append(code_block("""mvn test
mvn test -DsuiteXmlFile=testng-user-management.xml
mvn test -DsuiteXmlFile=testng-login-ui-ux.xml"""))
    story.append(Paragraph("<b>Maven cache issue:</b> If Maven reports AccessDeniedException under .m2, close Eclipse/Java processes and remove the affected plugin folder only after verifying the path.", note_style))

    story.append(h1("10. Dynamic Workflow and Role-Based Testing"))
    story.append(p("Some EasyQ workflows are dynamic: a module may show existing data, no data, restricted access, draft state, review state or approved state. The framework checks valid states instead of directly skipping workflows."))
    story.extend(bullet_list([
        "DynamicWorkflowHelper checks for meaningful page states such as Draft, Review, Approved, Pending, Completed, Closed, Open, No Data or Access Denied.",
        "Mutation-heavy tests are controlled by allowWorkflowMutations or EASYQ_ALLOW_WORKFLOW_MUTATIONS.",
        "Role-based tests use configured Admin, Document Controller and Assignee accounts.",
        "Disposable user tests use EASYQ_TEST_USER_* variables and include cleanup behavior.",
    ]))
    story.append(rl_table(["Role Type", "Username Key", "Password Key"], [
        ("Admin", "EASYQ_ADMIN_USERNAME", "EASYQ_ADMIN_PASSWORD or EASYQ_PASSWORD"),
        ("Document Controller", "EASYQ_DOC_CONTROLLER_USERNAME", "EASYQ_DOC_CONTROLLER_PASSWORD"),
        ("Assignee", "EASYQ_ASSIGNEE_SWATI_USERNAME", "EASYQ_ASSIGNEE_SWATI_PASSWORD"),
        ("Other Assignees", "EASYQ_ASSIGNEE_*_USERNAME", "matching password variable if used"),
    ], [1.45, 2.25, 2.65], font_size=7.8))

    story.append(h1("11. Reporting and Debugging"))
    story.append(p("TestListener prints failed test name, test data, failure reason and screenshot path. Several tests also print navigation or login step logs to make stuck points visible."))
    story.append(code_block("""FAILED TEST: verifyUserManagementPageLoadsSuccessfully
FAILURE REASON: <root cause>
SCREENSHOT: C:\\...\\test-output\\screenshots\\verifyUserManagementPageLoadsSuccessfully.png"""))
    story.append(h2("Common Issues and Fixes"))
    story.append(rl_table(["Issue", "Meaning", "Fix"], [
        ("EASYQ_PASSWORD is required", "Password variable is missing.", "Set EASYQ_ADMIN_PASSWORD or EASYQ_PASSWORD in Eclipse."),
        ("Login stays on login page", "Password may be wrong, account inactive, or field/click issue.", "Verify credentials manually, then rerun."),
        ("CDP implementation warning", "Chrome is newer than Selenium devtools package.", "Usually non-blocking unless DevTools APIs are used."),
        ("NoSuchWindowException", "Chrome was manually closed or crashed.", "Rerun; teardown is now safer and next test starts fresh."),
        ("Timeout waiting for element", "Page, module or control did not render in time.", "Use module suite, confirm navigation, increase visual delay/wait."),
        ("Unsupported class file major version", "Java/Maven/TestNG runtime mismatch.", "Use JDK 17 consistently in Eclipse and Maven."),
        (".m2 AccessDeniedException", "Maven cache folder locked or permission issue.", "Close Java/Eclipse processes; remove only affected plugin cache if it exists."),
    ], [1.5, 2.3, 2.55], font_size=7.2))

    story.append(h1("12. GitHub Workflow"))
    story.append(p("The project is connected to GitHub with origin remote configured. Use small commits after stable changes, and avoid committing secrets.local.properties, screenshots containing sensitive data, or real passwords."))
    story.append(rl_table(["Item", "Value"], [
        ("Remote", "https://github.com/SAURABH188/easyq-automation..git"),
        ("Current branch", "master"),
        ("Recommended commit scope", "One framework/module change at a time"),
    ], [1.6, 4.75], font_size=8))
    story.append(h2("Manual Push Steps from Eclipse"))
    story.extend(number_list([
        "Right-click project > Team > Commit.",
        "Review changed files and enter a clear commit message.",
        "Click Commit and Push, or Commit first and push later.",
        "If Push Ref Specification appears empty, choose Add All Branches Spec or select master -> master.",
        "Confirm Push Results shows master -> master or your current branch pushed successfully.",
    ]))
    story.append(h2("Command Line Backup"))
    story.append(code_block("""git status
git add .
git commit -m \"Update EasyQ automation framework waits and documentation\"
git push origin master"""))

    story.append(h1("13. Demo Talking Points"))
    story.extend(number_list([
        "Start by explaining the purpose: automate EasyQ beta validation across login, dashboard and QMS modules.",
        "Show pom.xml and explain Java, Selenium, TestNG and WebDriverManager.",
        "Show config.properties and explain that passwords are outside source code.",
        "Show testng XML files and explain module-wise execution.",
        "Open one test class and point out @Test description and manual test case ID comments.",
        "Show BaseTest, ActionHelper and WaitHelper to explain stability handling.",
        "Run a small suite such as login UI/UX or User Management page load.",
        "Show console output and screenshot behavior if a test fails.",
        "Close with next improvements: stable locators, CI, reports and cleaner page objects.",
    ]))

    story.append(h1("14. Maintenance Checklist"))
    story.append(rl_table(["Checklist Item", "Why It Matters"], [
        ("Keep credentials outside GitHub.", "Prevents password leakage."),
        ("Use module XML while debugging.", "Faster feedback than running the full suite."),
        ("Prefer stable locators.", "Reduces failures caused by UI layout changes."),
        ("Add manual test case ID comments.", "Keeps automation linked with QA test cases."),
        ("Capture useful assertion messages.", "Makes failed tests understandable without reopening code."),
        ("Use dynamic state handling for workflows.", "Avoids false failures when application data state changes."),
        ("Review screenshots after failures.", "Fastest way to confirm script, data, role or app-loading issue."),
        ("Commit after a passing compile.", "Keeps GitHub history clean and demo-safe."),
    ], [2.1, 4.25], font_size=7.8))

    story.append(h1("Appendix A. Useful Commands"))
    story.append(code_block("""# Run full suite if Maven cache is healthy
mvn test

# Run a module suite
mvn test -DsuiteXmlFile=testng-user-management.xml

# Check Git status
git status

# Push current branch
git push origin master

# Useful Eclipse run environment examples
EASYQ_ADMIN_PASSWORD=<password>
EASYQ_VISUAL_DELAY_MS=1500
EASYQ_PAGE_SETTLE_MS=1200"""))

    story.append(h1("Appendix B. Current Suite Map"))
    suite_rows = []
    for xml in sorted(ROOT.glob("testng*.xml")):
        text = xml.read_text(encoding="utf-8", errors="ignore")
        suite = re.search(r'<suite\s+name="([^"]+)"', text)
        classes = ", ".join(re.findall(r'<class\s+name="tests\.([^"]+)"', text)) or "-"
        suite_rows.append((xml.name, suite.group(1) if suite else "-", classes))
    story.append(rl_table(["Suite File", "Suite Name", "Classes"], suite_rows, [1.55, 2.15, 2.65], font_size=6.8))

    story.append(h1("Appendix C. Recommended Next Improvements"))
    story.extend(bullet_list([
        "Ask developers to add data-testid attributes to key buttons, inputs, menus and module links.",
        "Move remaining module classes toward a common BaseTest and page-object structure.",
        "Add ExtentReports or Allure for richer demo reporting.",
        "Create CI execution for nightly smoke suites once credentials are stored in CI secrets.",
        "Keep one stable smoke suite separate from long workflow mutation tests.",
    ]))

    pdf_doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="EasyQ Automation Project Guide",
        author="Codex for Saurabh",
    )
    pdf_doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return pdf_path


pdf_out = build_pdf()
print(OUT)
print(pdf_out)
